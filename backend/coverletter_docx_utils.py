import io
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import subprocess


def build_coverletter_docx(cover_letter: dict) -> bytes:
    """
    Builds a .docx cover letter from a CoverLetter-schema dict and returns the
    document as raw bytes.

    Uses paragraph `space_after` instead of blank spacer paragraphs between
    sections — this reclaims vertical space that was previously wasted on
    empty lines, helping keep standard-length letters on a single page.
    """
    header = cover_letter.get("header", {})
    employer = cover_letter.get("employers_info", {})

    doc = Document()

    # --- Base document styling ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    def add_para(text, space_after=8, bold=False, size=None, justify=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(space_after)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # --- Candidate header block ---
    if header.get("full_name"):
        add_para(header["full_name"], bold=True, size=14, space_after=2)

    contact_bits = [header.get("email"), header.get("phone"), header.get("address")]
    contact_line = " | ".join(b for b in contact_bits if b)
    if contact_line:
        add_para(contact_line, space_after=2)

    if header.get("date"):
        add_para(header["date"], space_after=14)

    # --- Employer info block ---
    employer_lines = [
        employer.get("hiring_manager_name"),
        employer.get("hiring_manager_title"),
        employer.get("company_name"),
        employer.get("company_address"),
    ]
    employer_lines = [l for l in employer_lines if l]
    for i, line in enumerate(employer_lines):
        last = i == len(employer_lines) - 1
        add_para(line, space_after=14 if last else 2)

    # --- Salutation ---
    if cover_letter.get("salutation"):
        add_para(cover_letter["salutation"], space_after=10)

    # --- Opening paragraph ---
    if cover_letter.get("opening_paragraph"):
        add_para(cover_letter["opening_paragraph"], justify=True, space_after=10)

    # --- Body paragraphs ---
    for bp in cover_letter.get("body_paragraphs", []):
        content = bp.get("content") if isinstance(bp, dict) else bp
        if content:
            add_para(content, justify=True, space_after=10)

    # --- Closing paragraph ---
    if cover_letter.get("closing_paragraph"):
        add_para(cover_letter["closing_paragraph"], justify=True, space_after=14)

    # --- Sign-off ---
    if cover_letter.get("sign_off"):
        add_para(cover_letter["sign_off"], space_after=2)

    if cover_letter.get("signature_name"):
        add_para(cover_letter["signature_name"], space_after=0)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

import subprocess
from pathlib import Path


def get_docx_page_count(docx_path: str) -> int:
    """
    Converts a .docx to PDF via headless LibreOffice and returns the page
    count using `pdfinfo`. Cleans up the intermediate PDF and temp directory
    afterward, regardless of success or failure.
    """
    docx_path = Path(docx_path)
    out_dir = docx_path.parent / f".pagecheck_{docx_path.stem}"
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        subprocess.run(
            [
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", str(out_dir), str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=60,
        )

        pdf_path = out_dir / f"{docx_path.stem}.pdf"
        result = subprocess.run(
            ["pdfinfo", str(pdf_path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=30,
        )

        for line in result.stdout.splitlines():
            if line.startswith("Pages:"):
                return int(line.split(":")[1].strip())

        raise RuntimeError(f"Could not determine page count for {pdf_path}")

    finally:
        shutil.rmtree(out_dir, ignore_errors=True)