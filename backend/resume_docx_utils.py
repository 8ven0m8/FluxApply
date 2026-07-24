"""
LangGraph node: generate_resume_node

Looks at the styling of the user's existing/original resume (font, sizes,
accent color, margins, whether section headers have an underline rule,
whether bullets are used) and re-uses that same visual design to render a
brand-new .docx populated with the tailored_resume content for a given JD.

Accepts the original resume as either a .docx or a .pdf (path or raw bytes).
PDFs are transparently converted to .docx (via pdf2docx) before style
extraction runs, so the rest of the pipeline never has to know which format
the user actually uploaded. Requires: pip install pdf2docx --break-system-packages

Nothing here is a separate importable module you need to maintain elsewhere —
drop this whole file in as your node file (or paste the functions straight
into your existing node file) and wire generate_resume_node into your graph.
"""

import os
import tempfile
from collections import Counter
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ----------------------------------------------------------------------
# 1. STYLE EXTRACTION — read an existing resume, infer its design system
# ----------------------------------------------------------------------

DEFAULT_STYLE = {
    "body_font": "Calibri",
    "body_size": 10.5,
    "name_font": "Calibri",
    "name_size": 22.0,
    "heading_size": 12.0,
    "accent_color": "1F4E5F",
    "text_color": "404040",
    "use_heading_rule": True,
    "use_bullets": True,
    "margin_in": 0.6,
}


def _rgb_to_hex(color):
    if color is None or color.type is None:
        return None
    try:
        return str(color.rgb)
    except Exception:
        return None


def _has_bottom_border(paragraph):
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        return False
    return pBdr.find(qn("w:bottom")) is not None


def _looks_like_pdf(source) -> bool:
    """
    source can be a filesystem path (str) or raw file bytes. Detect PDF
    either by the %PDF- magic bytes or the .pdf extension, so callers don't
    need to tell us the type up front.
    """
    if isinstance(source, (bytes, bytearray)):
        return source[:5] == b"%PDF-"
    if isinstance(source, str):
        return source.lower().endswith(".pdf")
    return False


def _pdf_to_docx_bytes(source) -> bytes:
    """
    Convert a PDF (path or raw bytes) into .docx bytes so the existing
    docx-based style extraction can run on it unchanged.

    pdf2docx only operates on real files on disk, so we round-trip through
    a temp .pdf -> temp .docx and read the result back into memory.
    """
    from pdf2docx import Converter  # lazy import: optional dependency

    tmp_pdf_path = None
    tmp_docx_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
            if isinstance(source, (bytes, bytearray)):
                tmp_pdf.write(source)
            else:
                with open(source, "rb") as f:
                    tmp_pdf.write(f.read())
            tmp_pdf_path = tmp_pdf.name

        tmp_docx_path = tmp_pdf_path[:-4] + ".docx"

        converter = Converter(tmp_pdf_path)
        try:
            converter.convert(tmp_docx_path)
        finally:
            converter.close()

        with open(tmp_docx_path, "rb") as f:
            return f.read()
    finally:
        for path in (tmp_pdf_path, tmp_docx_path):
            if path and os.path.exists(path):
                os.remove(path)


def extract_style_profile(source) -> dict:
    """
    source: a path to a .docx/.pdf, or raw bytes of a .docx/.pdf (the user's
    existing/original resume, whatever design it uses).
    Returns a style dict compatible with build_resume_docx(). Falls back to
    DEFAULT_STYLE for anything it can't confidently detect (including a PDF
    that fails to convert), so a weird, scanned, or heavily-tabled resume
    degrades gracefully instead of throwing.
    """
    if _looks_like_pdf(source):
        try:
            source = _pdf_to_docx_bytes(source)
        except Exception:
            return dict(DEFAULT_STYLE)

    try:
        doc = Document(BytesIO(source) if isinstance(source, (bytes, bytearray)) else source)
    except Exception:
        return dict(DEFAULT_STYLE)

    style = dict(DEFAULT_STYLE)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return style

    try:
        sec = doc.sections[0]
        style["margin_in"] = round(sec.left_margin.inches, 2)
    except Exception:
        pass

    first_p = paragraphs[0]
    if first_p.runs:
        biggest = max(first_p.runs, key=lambda r: (r.font.size.pt if r.font.size else 0))
        if biggest.font.size:
            style["name_size"] = biggest.font.size.pt
        if biggest.font.name:
            style["name_font"] = biggest.font.name
        hex_color = _rgb_to_hex(biggest.font.color)
        if hex_color:
            style["accent_color"] = hex_color

    signatures = Counter()
    sig_examples = {}
    heading_rule_votes = []
    body_sizes = Counter()
    body_fonts = Counter()

    for p in paragraphs[1:]:
        for r in p.runs:
            if not r.text.strip():
                continue
            size = r.font.size.pt if r.font.size else None
            color = _rgb_to_hex(r.font.color)
            if r.bold and (color or size):
                sig = (size, color)
                signatures[sig] += 1
                sig_examples[sig] = p
            elif not r.bold:
                if size:
                    body_sizes[size] += 1
                if r.font.name:
                    body_fonts[r.font.name] += 1

    if signatures:
        (heading_size, heading_color), _ = signatures.most_common(1)[0]
        if heading_size:
            style["heading_size"] = heading_size
        if heading_color:
            style["accent_color"] = heading_color
        example_p = sig_examples[(heading_size, heading_color)]
        heading_rule_votes.append(_has_bottom_border(example_p))

    if heading_rule_votes:
        style["use_heading_rule"] = any(heading_rule_votes)

    if body_sizes:
        style["body_size"] = body_sizes.most_common(1)[0][0]
    if body_fonts:
        style["body_font"] = body_fonts.most_common(1)[0][0]

    style["use_bullets"] = any(
        p.style is not None and "List Bullet" in (p.style.name or "")
        for p in doc.paragraphs
    )

    return style



def _set_margins(doc, inches):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _add_bottom_border(paragraph, color, size=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_borders.append(bottom)
    p_pr.append(p_borders)


def _section_heading(doc, text, style):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(style["heading_size"])
    run.font.color.rgb = RGBColor.from_string(style["accent_color"])
    if style["use_heading_rule"]:
        _add_bottom_border(p, style["accent_color"])
    return p


def _hyperlink(paragraph, url, text, color="0563C1", underline=True, size_pt=None):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    if size_pt is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
        rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _label_from_url(url):
    lowered = url.lower()
    if "linkedin" in lowered:
        return "LinkedIn"
    if "github" in lowered:
        return "GitHub"
    if "kaggle" in lowered:
        return "Kaggle"
    if "huggingface" in lowered:
        return "HuggingFace"
    return url.split("//")[-1].split("/")[0]

def _right_tab_position(doc, style):
    section = doc.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin
    return usable


def _add_right_tab(paragraph, doc, style):
    """Adds a right-aligned tab stop at the page's usable-width edge, so a
    run added after '\t' lands flush against the right margin — used to put
    project/date info on the same line as a title without a second paragraph."""
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        _right_tab_position(doc, style), WD_TAB_ALIGNMENT.RIGHT
    )


def _append_links_inline(paragraph, links, style, size_delta=-1.5, prefix="  "):
    """
    Appends pipe-separated labeled hyperlinks to the END of an existing
    paragraph (same line), instead of starting a new paragraph/line. Used to
    keep project and achievement links from costing an extra line each.
    """
    if not links:
        return
    size_pt = style["body_size"] + size_delta
    first = True
    for link in links:
        if not isinstance(link, dict):
            continue
        url = link.get("url")
        if not url:
            continue
        label = link.get("label") or _label_from_url(url)
        if first:
            sep_run = paragraph.add_run(prefix)
            sep_run.font.size = Pt(size_pt)
        else:
            sep_run = paragraph.add_run("  |  ")
            sep_run.font.size = Pt(size_pt)
        _hyperlink(paragraph, url, label, size_pt=size_pt)
        first = False


def _ensure_list(value):
    if value is None:
        return []
    if isinstance(value, str):
        return [value]
    if isinstance(value, (list, tuple)):
        return list(value)
    return [str(value)]


def build_resume_docx(tailored_resume: dict, style: dict = None) -> bytes:
    style = style or DEFAULT_STYLE
    details = tailored_resume.get("details", {})
    skills = tailored_resume.get("skills", [])
    summary = tailored_resume.get("summary", "")
    experience = tailored_resume.get("experience", [])
    projects = tailored_resume.get("projects", [])
    achievements = tailored_resume.get("achievements", [])
    publications = tailored_resume.get("publications", [])

    doc = Document()
    _set_margins(doc, style["margin_in"])

    normal = doc.styles["Normal"]
    normal.font.name = style["body_font"]
    normal.font.size = Pt(style["body_size"])
    normal.font.color.rgb = RGBColor.from_string(style["text_color"])

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(details.get("full_name", ""))
    name_run.bold = True
    name_run.font.name = style["name_font"]
    name_run.font.size = Pt(style["name_size"])
    name_run.font.color.rgb = RGBColor.from_string(style["accent_color"])

    # Contact
    contact_bits = [b for b in [details.get("email"), details.get("phone"), details.get("address")] if b]
    if contact_bits:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(2)
        contact_p.add_run("  |  ".join(contact_bits)).font.size = Pt(10)

    # Links
    profile_urls = details.get("profile_urls", [])
    if profile_urls:
        links_p = doc.add_paragraph()
        links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_p.paragraph_format.space_after = Pt(6)
        for i, url in enumerate(profile_urls):
            if i > 0:
                links_p.add_run("   |   ").font.size = Pt(10)
            _hyperlink(links_p, url, _label_from_url(url))

    if summary:
        _section_heading(doc, "Summary", style)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.add_run(summary).font.size = Pt(style["body_size"])

    education = details.get("education", [])
    institutions = details.get("institutions", [])
    if education or institutions:
        _section_heading(doc, "Education", style)
        for i, edu in enumerate(education):
            inst = institutions[i] if i < len(institutions) else (institutions[0] if institutions else "")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(inst)
            run.bold = True
            run.font.size = Pt(style["body_size"])
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(6)
            p2.add_run(edu).font.size = Pt(style["body_size"])
            
    if skills:
        _section_heading(doc, "Skills", style)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.add_run(" • ".join(skills)).font.size = Pt(style["body_size"])

    if experience:
        _section_heading(doc, "Experience", style)
        for exp in experience:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_after = Pt(0)
            role = exp.get("role") or exp.get("title", "")
            company = exp.get("company", "")
            run = title_p.add_run(f"{role} — {company}" if company else role)
            run.bold = True
            run.font.size = Pt(style["body_size"] + 0.5)
            dates = exp.get("dates") or exp.get("duration")
            if dates:
                d_run = title_p.add_run(f"\t{dates}")
                d_run.italic = True
                d_run.font.size = Pt(style["body_size"] - 1)

            bullets = _ensure_list(
                exp.get("bullets") if exp.get("bullets") is not None
                else exp.get("description")
            )
            for bullet in bullets:
                bp = doc.add_paragraph(style="List Bullet" if style["use_bullets"] else "Normal")
                bp.paragraph_format.space_after = Pt(2)
                bp.add_run(("• " if not style["use_bullets"] else "") + bullet).font.size = Pt(style["body_size"])

    if projects:
        _section_heading(doc, "Projects", style)
        for proj in projects:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_after = Pt(1)
            run = title_p.add_run(proj.get("title", ""))
            run.bold = True
            run.font.size = Pt(style["body_size"] + 0.5)

            links = proj.get("links", [])
            if links:
                _add_right_tab(title_p, doc, style)
                title_p.add_run("\t")
                _append_links_inline(title_p, links, style, prefix="")

            tech = proj.get("technologies", [])
            if tech:
                tech_p = doc.add_paragraph()
                tech_p.paragraph_format.space_after = Pt(1)
                tech_run = tech_p.add_run(" | ".join(tech))
                tech_run.italic = True
                tech_run.font.size = Pt(style["body_size"] - 1)
                tech_run.font.color.rgb = RGBColor.from_string(style["accent_color"])

            desc = proj.get("description", "")
            if desc:
                desc_p = doc.add_paragraph()
                desc_p.paragraph_format.space_after = Pt(6)
                desc_p.add_run(desc).font.size = Pt(style["body_size"])

    if achievements:
        _section_heading(doc, "Achievements", style)
        for ach in achievements:
            # Backward compatible: older stored data may have plain strings
            # instead of {"text": ..., "links": [...]}.
            if isinstance(ach, dict):
                text = ach.get("text", "")
                links = ach.get("links", [])
            else:
                text, links = str(ach), []

            bp = doc.add_paragraph(style="List Bullet" if style["use_bullets"] else "Normal")
            bp.paragraph_format.space_after = Pt(2)
            run = bp.add_run(("• " if not style["use_bullets"] else "") + text)
            run.font.size = Pt(style["body_size"])
            _append_links_inline(bp, links, style)

    if publications:
        _section_heading(doc, "Publications", style)
        for pub in publications:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_after = Pt(1)
            run = title_p.add_run(pub.get("title", ""))
            run.bold = True
            run.font.size = Pt(style["body_size"] + 0.5)

            meta_bits = [b for b in [pub.get("venue"), pub.get("date"), pub.get("authors")] if b]
            if meta_bits:
                _add_right_tab(title_p, doc, style)
                title_p.add_run("\t")
                meta_run = title_p.add_run(" | ".join(meta_bits))
                meta_run.italic = True
                meta_run.font.size = Pt(style["body_size"] - 1)

            if pub.get("link"):
                link_p = doc.add_paragraph()
                link_p.paragraph_format.space_after = Pt(1)
                _append_links_inline(link_p, [{"label": "Link", "url": pub["link"]}], style, prefix="")

            if pub.get("description"):
                desc_p = doc.add_paragraph()
                desc_p.paragraph_format.space_after = Pt(6)
                desc_p.add_run(pub["description"]).font.size = Pt(style["body_size"])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()