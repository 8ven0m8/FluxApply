#----------------------------------------------------#
# Extracts every important detail from the resume    #
# and provides in a beautiful json format using llms #
#----------------------------------------------------#

from os import getenv
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI 
from playwright.async_api import async_playwright
from typing import TypedDict, Optional, List
from pydantic import BaseModel, Field
from langgraph.store.base import BaseStore
from langchain_core.runnables import RunnableConfig
from langgraph.graph import StateGraph, START, END
from langchain_core.output_parsers import PydanticOutputParser
from pypdf import PdfReader
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
from schemas import Details, Experience, Project, ResumeFact
from system_prompts import SYSTEM_PROMPT_FOR_RESUME_EXTRACTION
from usage_tracking import log_llm_usage

load_dotenv()
llm = ChatOpenAI(
    # base_url=getenv("FREELLMAPI_URL"),
    api_key=getenv("OPENAI_KEY"),
    model="gpt-4o-mini"
)


############### Extract text from pdf or doc #################
# text extraction
def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    text_parts = []
    for page in reader.pages:
        text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)
def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)
def extract_resume_text(path: str) -> str:
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only .pdf and .docx supported.")
# link extraction
def extract_links_from_pdf(path: str) -> list[str]:
    reader = PdfReader(path)
    links = []
    for page in reader.pages:
        if "/Annots" in page:
            for annot in page["/Annots"]:
                obj = annot.get_object()
                if "/A" in obj and "/URI" in obj["/A"]:
                    links.append(obj["/A"]["/URI"])
    return links
def extract_links_from_docx(path: str) -> list[str]:
    doc = Document(path)
    links = []

    rels = doc.part.rels

    for paragraph in doc.paragraphs:
        for hyperlink in paragraph._p.findall(qn('w:hyperlink')):
            rel_id = hyperlink.get(qn('r:id'))
            if rel_id and rel_id in rels:
                url = rels[rel_id].target_ref
                if url not in links:
                    links.append(url)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for hyperlink in paragraph._p.findall(qn('w:hyperlink')):
                        rel_id = hyperlink.get(qn('r:id'))
                        if rel_id and rel_id in rels:
                            url = rels[rel_id].target_ref
                            if url not in links:
                                links.append(url)

    return links
def extract_links(path: str) -> list[str]:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return extract_links_from_pdf(path)
    elif ext == ".docx":
        return extract_links_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

############################# States ################################
class ResumeState(TypedDict):
    resume_text: Optional[str]
    resume_facts: Optional[ResumeFact]

resume_extraction_parser = PydanticOutputParser(
    pydantic_object=ResumeFact
)

####################### Node ######################
def resume_extraction(state: ResumeState, config: RunnableConfig, *, store: BaseStore):
    resume_text = state.get("resume_text")
    if not resume_text:
        raise ValueError("No resume_text found in state, nothing to extract from.")

    user_id = config["configurable"]["user_id"]

    prompt = (
        f"{resume_extraction_parser.get_format_instructions()}\n\n"
        f"{SYSTEM_PROMPT_FOR_RESUME_EXTRACTION}\n\n"
        f"RESUME TEXT:\n{resume_text}"
    )

    output = llm.invoke(prompt)
    log_llm_usage(user_id=user_id, endpoint="resume_upload", node_name="resume_extraction", ai_message=output)
    facts: ResumeFact = resume_extraction_parser.parse(output.content)

    ns = ("user", user_id, "resume_facts")
    store.put(ns, "current", {"data": facts.model_dump()})

    return {"resume_facts": facts}

###################### Build graph #####################
builder = StateGraph(ResumeState)
builder.add_node('extract_resume', resume_extraction)

builder.add_edge(START, 'extract_resume')
builder.add_edge('extract_resume', END)

DB_URI = getenv("DB_URI")

def process_resume(file_path: str, user_id: str, store: BaseStore) -> str:
    """
    `store` is the single pooled PostgresStore created once at app startup
    and injected by the caller (see app.py's `get_store` dependency) — this
    function no longer opens its own Postgres connection per call.
    """
    resume_text = extract_resume_text(file_path)
    links = extract_links(file_path)

    if links:
        resume_text += "\n\nPROFILE LINKS:\n" + "\n".join(links)

    graph = builder.compile(store=store)
    config = {"configurable": {"user_id": user_id}}
    out = graph.invoke({"resume_text": resume_text}, config)

    return out["resume_facts"].model_dump_json(indent=2)